from django.shortcuts import render,get_object_or_404
from django.views.generic import ListView,View,DetailView,UpdateView,DeleteView
from .models import Customers,Account,Card,Currencys
from django.shortcuts import redirect
from django.http.response import HttpResponseRedirect
from django.urls import reverse_lazy
from datetime import datetime
from django.http import FileResponse
from django.views.generic import View
from docx import Document
from datetime import datetime
import io


class IndexViews(ListView):
    template_name = "index.html"
    model = Customers



class AddCustomerView(ListView):
    template_name = "add_customers.html"
    model = Customers

    def post(self, request, **kwargs):
        last_name = request.POST.get('last_name')
        first_name = request.POST.get('first_name')
        middle_name = request.POST.get('middle_name')
        date_of_birth = request.POST.get('date_of_birth')
        gender = request.POST.get('gender')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        address = request.POST.get('address')
        passport_number = request.POST.get('passport_number')
        password_issue_data = request.POST.get('password_issue_data')
        photo = request.FILES.get('photo')

        check_passport_number=Customers.get_customer_by_paspord(passport_number)

        if not check_passport_number:
            customer=Customers(last_name = last_name,
                               first_name = first_name,
                               middle_name = middle_name,
                               gender = gender,
                               address = address,
                               email=email,
                               phone_numder = phone,
                               password_number = passport_number,
                               password_issue_data = password_issue_data,
                               date_of_birth=date_of_birth,
                               photo=photo
                               )
            customer.save()
            # request.session["customer"]=email


            response=HttpResponseRedirect(redirect_to="/")
            return response
        else:
            response=HttpResponseRedirect(redirect_to=".")
            return response


class CustomerDetailView(DetailView):
    model = Customers
    template_name = 'customers_detail.html'
    context_object_name = 'customer'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        customer = self.get_object()
        context['accounts'] =Account.objects.filter(customer=customer)
        return context

    def post(self, request, **kwargs):
        account_pk = request.POST.get('account_pk')
        account = get_object_or_404(Account, id=account_pk)
        account.status = "disactive"
        account.closed_at   = datetime.now()

        account.save()
        return self.get(request)



class AccountCardsListView(ListView):
    model = Card
    template_name = 'account_cards.html'
    context_object_name = 'cards'

    def get_queryset(self):
        account_id = self.kwargs['pk']
        account = get_object_or_404(Account, pk=account_id)
        return Card.objects.filter(account=account)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        account_id = self.kwargs['pk']
        context['account'] = get_object_or_404(Account, pk=account_id)
        return context
    def post(self, request, **kwargs):
        account_pk = request.POST.get('account_pk')
        customer_pk = request.POST.get('customer_pk')
        prefix = '2200'  # Префикс для карт "Мир"
        length = 16  # Длина номера карты
        valid_card_number = generate_valid_card_number(prefix, length)
        current_time=datetime.now()
        current_time=add_years(current_time,10)





        card=Card(
            card_number=valid_card_number,
            expiration_date=current_time,
            customer=Customers.objects.get(pk=customer_pk),
            account=Account.objects.get(pk=account_pk)
        )
        card.save()
        # response=HttpResponseRedirect(redirect_to = '')
        return self.get(request)



class EditCustomerView(UpdateView):
    model = Customers
    fields = []  # Список полей, которые будут доступны для редактирования
    template_name = 'edit_customer.html'
    success_url = '/customer_list/'  # URL для перенаправления после успешного редактирования

    def get_object(self, queryset=None):
        # Получаем клиента по его идентификатору
        return get_object_or_404(Customers, id=self.kwargs['pk'])

    def post(self, request, **kwargs):
        customer = get_object_or_404(Customers, id=self.kwargs['pk'])
        last_name = request.POST.get('last_name')
        first_name = request.POST.get('first_name')
        middle_name = request.POST.get('middle_name')
        date_of_birth = request.POST.get('date_of_birth')
        gender = request.POST.get('gender')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        address = request.POST.get('address')
        password_number = request.POST.get('passport_number')
        password_issue_data = request.POST.get('password_issue_data')

        # check_passport_number=Customers.get_customer_by_paspord(password_number)

        customer.last_name = last_name
        customer.first_name = first_name
        customer.middle_name = middle_name
        customer.gender = gender
        customer.address = address
        customer.email = email
        customer.phone_numder = phone
        customer.password_number = password_number
        customer.password_issue_data = password_issue_data
        customer.date_of_birth = date_of_birth



        customer.save()
        # request.session["customer"]=email


        return redirect(reverse_lazy('customer_detail', kwargs={'pk': self.kwargs['pk']}))
class CustomerDeleteView(DeleteView):
    model = Customers
    template_name = 'confirm_delete_customer.html'
    success_url = reverse_lazy('customer_list')

class AddAccountView(ListView):
    model = Account
    template_name = "add_account.html"
    def get(self, request,pk):

        currencies = Currencys.objects.all()
        return render(request, self.template_name, {'currencies': currencies ,"customer":pk})

    def post(self, request, **kwargs):
        customer_pk = request.POST.get("customer_pk")
        type_account =request.POST.get("type_account")
        currency_id =request.POST.get("currency")



        account =Account(customer=Customers.objects.get(pk=customer_pk),
                         account_type=type_account,
                         currency=Currencys.objects.get(pk=currency_id),
                         opened_at=datetime.now())
        account.save()
        return redirect(reverse_lazy('customer_detail', kwargs={'pk': customer_pk}))

class CreateContractView(View):
    def get(self, request):
        # Получаем данные из параметров запроса
        customer_name = request.GET.get('customer_name', '')  # Получаем имя клиента из параметра запроса
        account_number = request.GET.get('account_number', '')  # Получаем номер счета из параметра запроса
        print(customer_name)
        # Создаем документ Word
        document = Document()
        document.add_heading('Договор о открытии счета', level=1)
        document.add_paragraph('Дата: ' + datetime.now().strftime('%d.%m.%Y'))
        document.add_paragraph()
        document.add_paragraph('Мы, банк "Название банка", именуемый в дальнейшем "Банк",')
        document.add_paragraph('с одной стороны, и гражданин ' + customer_name + ',')
        document.add_paragraph('действующий от своего имени, именуемый в дальнейшем "Клиент", с другой стороны,')
        document.add_paragraph('заключили настоящий договор о нижеследующем:')
        document.add_paragraph()
        document.add_paragraph('1. Банк открывает Клиенту счет под номером ' + account_number + '.')
        document.add_paragraph()
        document.add_paragraph('Настоящий договор вступает в силу с момента его подписания и действует в течение указанного периода.')
        document.add_paragraph()
        document.add_paragraph('Подпись Банка: _________________     Подпись Клиента: _________________')

        # Сохраняем документ в байтовый поток
        document_stream = io.BytesIO()
        document.save(document_stream)
        document_stream.seek(0)

        # Отправляем документ как файл в ответ на запрос
        response = FileResponse(document_stream, as_attachment=True, filename='открытие_счета_{}.docx'.format(customer_name))
        return response

def luhn_checksum(card_number):
    # Преобразуем номер карты в список цифр
    digits = [int(x) for x in str(card_number)]
    # Удваиваем каждую вторую цифру, начиная с предпоследней
    for i in range(len(digits) - 2, -1, -2):
        digits[i] *= 2
        # Если результат удвоения больше 9, вычитаем 9
        if digits[i] > 9:
            digits[i] -= 9
    # Суммируем все цифры
    checksum = sum(digits)
    # Проверяем, является ли сумма кратной 10
    return checksum % 10 == 0

def generate_valid_card_number(prefix, length):
    # Генерируем случайные цифры для заполнения оставшихся позиций
    import random
    random_digits = [random.randint(0, 9) for _ in range(length - len(prefix) - 1)]
    # Добавляем префикс к сгенерированным цифрам
    card_number = [int(x) for x in str(prefix)] + random_digits
    # Вычисляем контрольную сумму для валидации номера карты
    for i in range(10):
        card_number.append(i)
        if luhn_checksum(''.join(map(str, card_number))):
            return ''.join(map(str, card_number))
        card_number.pop()
    raise ValueError("Couldn't generate a valid card number")
def add_years(start_date, years):
    try:
        return start_date.replace(year=start_date.year + years)
    except ValueError:
        # 👇️ Preserve calendar day (if Feb 29th doesn't exist, set to 28th)
        return start_date.replace(year=start_date.year + years, day=28)

# Пример использования:


# Create your views here.
