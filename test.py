from docx import Document
from docx.shared import Pt
from datetime import datetime

def create_contract(customer_name, account_number):
    document = Document()

    # Заголовок договора
    document.add_heading('Договор о открытии счета', level=1)

    # Дата создания договора
    document.add_paragraph('Дата: ' + datetime.now().strftime('%d.%m.%Y'))
    document.add_paragraph()

    # Информация о клиенте
    document.add_paragraph('Мы, банк "Название банка", именуемый в дальнейшем "Банк",')
    document.add_paragraph('с одной стороны, и гражданин ' + "customer_name" + ',')
    document.add_paragraph('действующий от своего имени, именуемый в дальнейшем "Клиент", с другой стороны,')
    document.add_paragraph('заключили настоящий договор о нижеследующем:')
    document.add_paragraph()

    # Информация о счете
    document.add_paragraph('1. Банк открывает Клиенту счет под номером ' + "account_number" + '.')
    document.add_paragraph()

    # Заключение
    document.add_paragraph('Настоящий договор вступает в силу с момента его подписания и действует в течение указанного периода.')
    document.add_paragraph()

    # Подпись
    document.add_paragraph('Подпись Банка: _________________     Подпись Клиента: _________________')

    # Сохраняем документ
    document.save('открытие_счета_{}.docx'.format(customer_name))

# Пример использования
create_contract('Иванов Иван', '1234567890')
